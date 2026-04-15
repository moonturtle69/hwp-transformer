"""
EBS 올림포스 HWP 자동 변환기 - Streamlit 버전
실행: streamlit run app_streamlit.py
"""

import os
import subprocess
import tempfile
import shutil
import warnings
from io import BytesIO

import streamlit as st
from bs4 import BeautifulSoup, XMLParsedAsHTMLWarning
from docx import Document
from docx.shared import Pt, RGBColor, Mm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

warnings.filterwarnings("ignore", category=XMLParsedAsHTMLWarning)

CIRCLE_NUMS = '❶❷❸❹❺❻❼❽❾❿⓫⓬⓭⓮⓯⓰⓱⓲⓳⓴'

# ── 파싱 ──────────────────────────────────────────────────────────────────────

def hwp_to_html(hwp_bytes):
    tmpdir = tempfile.mkdtemp()
    try:
        hwp_path = os.path.join(tmpdir, 'input.hwp')
        with open(hwp_path, 'wb') as f:
            f.write(hwp_bytes)
        result = subprocess.run(
            ['hwp5html', hwp_path, '--output', tmpdir],
            capture_output=True, text=True
        )
        if result.returncode != 0:
            raise RuntimeError(f"hwp5html 실패: {result.stderr}")
        with open(os.path.join(tmpdir, 'index.xhtml'), encoding='utf-8') as f:
            return f.read()
    finally:
        shutil.rmtree(tmpdir, ignore_errors=True)


def extract_numbered_paras(td):
    result, current_num, current_parts = [], None, []
    for p in td.find_all('p'):
        text = p.get_text().replace('\r', '').strip()
        if not text:
            continue
        if text[0] in CIRCLE_NUMS:
            if current_num is not None:
                result.append((current_num, ' '.join(current_parts).strip()))
            current_num, current_parts = text[0], [text[1:].strip()] if text[1:].strip() else []
        elif current_num and text.strip():
            current_parts.append(text.strip())
    if current_num:
        result.append((current_num, ' '.join(current_parts).strip()))
    return result


def extract_question_spans(para):
    return [(s.get_text().replace('\r', ''), 'charshape-14' in (s.get('class') or []))
            for s in para.find_all('span')]


def parse_hwp_pages(html):
    soup = BeautifulSoup(html, 'lxml')
    pages = []
    for tbl in soup.find_all('table', class_='TableControl'):
        first_row = tbl.find('tr')
        if not first_row or '▶' not in first_row.get_text():
            continue
        rows = tbl.find_all('tr', recursive=False)
        if len(rows) < 2:
            continue

        category = rows[0].find('td').get_text().replace('\r', '').strip()
        q_td = rows[1].find('td')
        question_spans = []
        question_raw = ''

        first_p = True
        for p in q_td.find_all('p'):
            txt = p.get_text().replace('\r', '').strip()
            if not txt:
                continue
            cls = p.get('class', [])
            if first_p and not any(c.startswith('parashape') for c in cls):
                question_spans = extract_question_spans(p)
                question_raw = txt
                first_p = False
            else:
                first_p = False

        english_pairs = extract_numbered_paras(q_td)

        if len(rows) == 3:
            korean_pairs = extract_numbered_paras(rows[2].find('td'))
            pages.append({
                'needs_transform': True,
                'category': category,
                'question_spans': question_spans,
                'question_raw': question_raw,
                'english_pairs': english_pairs,
                'korean_pairs': korean_pairs,
            })
        else:
            pages.append({
                'needs_transform': False,
                'category': category,
                'question_spans': question_spans,
                'question_raw': question_raw,
                'english_pairs': english_pairs,
                'korean_pairs': [],
            })
    return pages


# ── DOCX 생성 ─────────────────────────────────────────────────────────────────

def add_run(para, text, font='맑은 고딕', size=10, bold=False, italic=False,
            underline=False, color=None):
    run = para.add_run(text)
    run.font.name = font
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.underline = underline
    if color:
        run.font.color.rgb = RGBColor(*color)
    rPr = run._r.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), font)
    return run


def set_table_border(table):
    tbl = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    tblBorders = OxmlElement('w:tblBorders')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        el = OxmlElement(f'w:{edge}')
        el.set(qn('w:val'), 'single')
        el.set(qn('w:sz'), '6')
        el.set(qn('w:space'), '0')
        el.set(qn('w:color'), '000000')
        tblBorders.append(el)
    tblPr.append(tblBorders)


def build_docx(pages):
    doc = Document()
    sec = doc.sections[0]
    sec.page_width = Mm(210)
    sec.page_height = Mm(297)
    sec.top_margin = Mm(15)
    sec.bottom_margin = Mm(15)
    sec.left_margin = Mm(16)
    sec.right_margin = Mm(16)

    for i, page in enumerate(pages):
        if i > 0:
            doc.add_page_break()

        # 헤더 테이블
        hdr = doc.add_table(rows=1, cols=3)
        hdr.style = 'Table Grid'
        c0, c1, c2 = hdr.rows[0].cells
        p1 = c1.paragraphs[0]
        p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_run(p1, 'EBS 올림포스 – 영어독해 기본 01', size=9, bold=True)
        p2 = c2.paragraphs[0]
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_run(p2, 'Unit 02 – 심경 (변화) · 분위기', size=8)
        doc.add_paragraph('')

        # 본문 테이블
        main = doc.add_table(rows=2, cols=1)
        set_table_border(main)

        # Row 0: 카테고리
        cat_p = main.rows[0].cells[0].paragraphs[0]
        cat_p.paragraph_format.space_before = Pt(2)
        cat_p.paragraph_format.space_after = Pt(2)
        add_run(cat_p, page.get('category', ''), bold=True)

        # Row 1: 내용
        cell = main.rows[1].cells[0]
        cell._tc.clear_content()

        # 문제 문장
        if page.get('question_spans'):
            q_p = cell.add_paragraph()
            q_p.paragraph_format.space_after = Pt(6)
            for text, special in page['question_spans']:
                if text:
                    add_run(q_p, text, bold=special, italic=special, underline=special)
        elif page.get('question_raw'):
            q_p = cell.add_paragraph()
            add_run(q_p, page['question_raw'])

        # 빈 줄
        sp = cell.add_paragraph('')
        sp.runs and setattr(sp.runs[0].font, 'size', Pt(4))

        # 한글 위 + 영문 아래 쌍 배치
        eng_dict = {c: t for c, t in page.get('english_pairs', [])}
        kor_dict = {c: t for c, t in page.get('korean_pairs', [])}
        circles = [c for c, _ in page.get('english_pairs', [])]

        for circle in circles:
            kor = kor_dict.get(circle, '')
            eng = eng_dict.get(circle, '')

            if kor:
                ko_p = cell.add_paragraph()
                ko_p.paragraph_format.space_before = Pt(2)
                ko_p.paragraph_format.space_after = Pt(0)
                add_run(ko_p, kor, size=9, color=(80, 80, 80))

            if eng:
                en_p = cell.add_paragraph()
                en_p.paragraph_format.space_before = Pt(0)
                en_p.paragraph_format.space_after = Pt(8)
                add_run(en_p, circle + ' ', bold=True)
                add_run(en_p, eng, font='Times New Roman')

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


# ── Streamlit UI ──────────────────────────────────────────────────────────────

st.set_page_config(
    page_title="EBS 올림포스 HWP 변환기",
    page_icon="📄",
    layout="centered"
)

st.title("📄 EBS 올림포스 HWP 자동 변환기")
st.caption("영어독해 기본 · 한글 번역을 영문 위에 자동 배치")

st.markdown("---")

# 변환 방식 설명
with st.expander("📖 변환 방식 보기", expanded=False):
    col1, col2 = st.columns(2)
    with col1:
        st.markdown("**변환 전 (원본)**")
        st.markdown("""
        ```
        ❶ Cooked food does many...
        ❷ It makes our food safer...
        ❸ Heating can allow us...

        [하단 한글 박스]
        ❶ 조리된 음식은 친숙한...
        ❷ 그것은 우리의 음식을...
        ❸ 가열은 우리가...
        ```
        """)
    with col2:
        st.markdown("**변환 후 (완성)**")
        st.markdown("""
        ```
        조리된 음식은 친숙한...
        ❶ Cooked food does many...

        그것은 우리의 음식을...
        ❷ It makes our food safer...

        가열은 우리가...
        ❸ Heating can allow us...
        ```
        """)
    st.info("✅ 번호 매칭 자동화 · ✅ 한글 번호 제거 · ✅ 하단 박스 삭제")

st.markdown("---")

# 파일 업로드
uploaded = st.file_uploader(
    "HWP 파일을 업로드하세요",
    type=['hwp'],
    help="변환 전 원본 HWP 파일 (최대 50MB)"
)

if uploaded:
    st.success(f"✅ 파일 업로드 완료: **{uploaded.name}**")

    col_preview, col_convert = st.columns(2)

    with col_preview:
        if st.button("🔍 미리보기", use_container_width=True):
            with st.spinner("파일 분석 중..."):
                try:
                    html = hwp_to_html(uploaded.read())
                    pages = parse_hwp_pages(html)
                    needs = sum(1 for p in pages if p['needs_transform'])

                    st.metric("전체 페이지", len(pages))
                    st.metric("변환 필요", needs)
                    st.metric("이미 완료", len(pages) - needs)

                    for i, p in enumerate(pages):
                        status = "⚠️ 변환 필요" if p['needs_transform'] else "✅ 완료"
                        st.write(f"**{i+1}페이지** — {p['category']} · {len(p.get('english_pairs', []))}문장 · {status}")
                except Exception as e:
                    st.error(f"오류: {e}")

    with col_convert:
        if st.button("✨ 변환 시작", type="primary", use_container_width=True):
            with st.spinner("변환 중..."):
                try:
                    uploaded.seek(0)
                    html = hwp_to_html(uploaded.read())
                    pages = parse_hwp_pages(html)

                    if not pages:
                        st.error("변환할 페이지를 찾을 수 없습니다.")
                    else:
                        buf = build_docx(pages)
                        base = os.path.splitext(uploaded.name)[0]
                        out_name = f"{base}_변환완료.docx"

                        st.success(f"✅ 변환 완료! {len(pages)}페이지 처리됨")
                        st.download_button(
                            label="⬇️ DOCX 다운로드",
                            data=buf,
                            file_name=out_name,
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            use_container_width=True
                        )
                except Exception as e:
                    st.error(f"변환 실패: {e}")

st.markdown("---")
st.caption("출력 형식: DOCX · 한글 프로그램에서 열고 HWP로 다시 저장 가능")
