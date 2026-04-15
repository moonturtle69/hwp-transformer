"""
HWP 영어독해 변환기
- 영문 문장 + 하단 한글 박스 → 한글을 영문 위에 배치, 한글 박스 삭제
"""

import os
import sys
import subprocess
import tempfile
import shutil
import warnings
import re
from io import BytesIO

from flask import Flask, request, send_file, render_template, jsonify
from bs4 import BeautifulSoup, XMLParsedAsHTMLWarning
from docx import Document
from docx.shared import Pt, RGBColor, Mm, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

warnings.filterwarnings("ignore", category=XMLParsedAsHTMLWarning)

app = Flask(__name__)
app.config['MAX_CONTENT_LENGTH'] = 50 * 1024 * 1024  # 50MB

CIRCLE_NUMS = '❶❷❸❹❺❻❼❽❾❿⓫⓬⓭⓮⓯⓰⓱⓲⓳⓴'

# ─────────────────────────────────────────
# HWP Parsing
# ─────────────────────────────────────────

def hwp_to_html(hwp_path):
    """Convert HWP file to HTML using hwp5html"""
    tmpdir = tempfile.mkdtemp()
    result = subprocess.run(
        ['hwp5html', hwp_path, '--output', tmpdir],
        capture_output=True, text=True
    )
    if result.returncode != 0:
        shutil.rmtree(tmpdir, ignore_errors=True)
        raise RuntimeError(f"hwp5html 변환 실패: {result.stderr}")
    html_path = os.path.join(tmpdir, 'index.xhtml')
    with open(html_path, 'r', encoding='utf-8') as f:
        content = f.read()
    shutil.rmtree(tmpdir, ignore_errors=True)
    return content


def extract_numbered_paras(td):
    """Extract numbered (❶❷❸...) paragraphs from a table cell."""
    paras = td.find_all('p')
    result = []
    current_num = None
    current_parts = []

    for p in paras:
        text = p.get_text().replace('\r', '').strip()
        if not text:
            continue
        if text[0] in CIRCLE_NUMS:
            if current_num is not None:
                result.append((current_num, ' '.join(current_parts).strip()))
            current_num = text[0]
            rest = text[1:].strip()
            current_parts = [rest] if rest else []
        elif current_num is not None:
            clean = text.strip()
            if clean:
                current_parts.append(clean)

    if current_num is not None:
        result.append((current_num, ' '.join(current_parts).strip()))

    return result


def extract_question_spans(para):
    """Return list of (text, is_bold) from question paragraph spans."""
    parts = []
    for span in para.find_all('span'):
        txt = span.get_text().replace('\r', '')
        classes = span.get('class', [])
        # charshape-14 is the bold/italic underlined English in question
        is_special = 'charshape-14' in classes
        parts.append((txt, is_special))
    return parts


def parse_hwp_pages(html):
    """
    Parse HWP HTML and return list of page data dicts.
    Each page with the ORIGINAL format (3 rows: header | English | Korean box)
    gets returned as a dict with 'needs_transform': True.
    Pages already in transformed format (2 rows) return 'needs_transform': False.
    """
    soup = BeautifulSoup(html, 'lxml')
    pages = []

    # Find content tables (those with ▶ in first row)
    for tbl in soup.find_all('table', class_='TableControl'):
        first_row = tbl.find('tr')
        if not first_row or '▶' not in first_row.get_text():
            continue

        rows = tbl.find_all('tr', recursive=False)
        if len(rows) < 2:
            continue

        # Category (▶ 함축 의미 추론 etc.)
        category = rows[0].find('td').get_text().replace('\r', '').strip()

        # Question row (always row index 1)
        q_td = rows[1].find('td')
        q_paras = q_td.find_all('p')

        # First non-empty paragraph = question text
        question_spans = []
        question_raw = ''
        english_paras = []

        first_p = True
        for p in q_paras:
            txt = p.get_text().replace('\r', '').strip()
            if not txt:
                continue
            cls = p.get('class', [])
            if first_p and not any(c.startswith('parashape') for c in cls):
                # This is the question paragraph
                question_spans = extract_question_spans(p)
                question_raw = txt
                first_p = False
            else:
                first_p = False

        english_pairs = extract_numbered_paras(q_td)

        if len(rows) == 3:
            # Original format: has Korean box at bottom
            kor_td = rows[2].find('td')
            korean_pairs = extract_numbered_paras(kor_td)
            pages.append({
                'needs_transform': True,
                'category': category,
                'question_spans': question_spans,
                'question_raw': question_raw,
                'english_pairs': english_pairs,
                'korean_pairs': korean_pairs,
            })
        else:
            # Already transformed (2 rows) - extract paired content
            # Parse alternating Korean / English lines from row 1
            paired = []
            paras = q_td.find_all('p')
            prev_kor = None
            for p in paras:
                txt = p.get_text().replace('\r', '').strip()
                if not txt:
                    continue
                cls = p.get('class', [])
                classes_str = ' '.join(cls)
                if txt[0] in CIRCLE_NUMS and 'parashape-22' in classes_str:
                    # English sentence
                    paired.append((prev_kor or '', txt[0], txt[1:].strip()))
                    prev_kor = None
                elif 'parashape-25' in classes_str and txt[0] not in CIRCLE_NUMS:
                    prev_kor = txt
            pages.append({
                'needs_transform': False,
                'category': category,
                'question_spans': question_spans,
                'question_raw': question_raw,
                'paired': paired,
            })

    return pages


# ─────────────────────────────────────────
# DOCX Generation
# ─────────────────────────────────────────

def set_cell_border(cell, **kwargs):
    """Set borders on a table cell."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        tag = f'w:{edge}'
        el = OxmlElement(tag)
        el.set(qn('w:val'), kwargs.get(edge, 'none'))
        el.set(qn('w:sz'), '4')
        el.set(qn('w:space'), '0')
        el.set(qn('w:color'), kwargs.get(f'{edge}_color', '000000'))
        tcBorders.append(el)
    tcPr.append(tcBorders)


def set_table_border(table, border_type='single', color='000000', size=4):
    """Set outer borders for a table."""
    tbl = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    tblBorders = OxmlElement('w:tblBorders')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        el = OxmlElement(f'w:{edge}')
        el.set(qn('w:val'), border_type)
        el.set(qn('w:sz'), str(size))
        el.set(qn('w:space'), '0')
        el.set(qn('w:color'), color)
        tblBorders.append(el)
    tblPr.append(tblBorders)


def add_run_with_style(para, text, font_name='맑은 고딕', size=10, bold=False,
                        color=None, underline=False, italic=False):
    run = para.add_run(text)
    run.font.name = font_name
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.underline = underline
    if color:
        run.font.color.rgb = RGBColor(*color)
    # Set East Asian font
    rPr = run._r.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), font_name)
    return run


def add_question_para(cell, question_spans):
    """Add the question paragraph with mixed styling."""
    para = cell.add_paragraph()
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.space_after = Pt(4)

    for text, is_special in question_spans:
        if not text:
            continue
        add_run_with_style(
            para, text,
            font_name='맑은 고딕',
            size=10,
            bold=is_special,
            underline=is_special,
            italic=is_special
        )
    return para


def add_korean_line(cell, kor_text):
    """Add Korean translation line (gray, smaller)."""
    para = cell.add_paragraph()
    para.paragraph_format.space_before = Pt(2)
    para.paragraph_format.space_after = Pt(0)
    pPr = para._p.get_or_add_pPr()
    # Indent slightly
    ind = OxmlElement('w:ind')
    ind.set(qn('w:left'), '0')
    pPr.append(ind)
    add_run_with_style(
        para, kor_text,
        font_name='맑은 고딕',
        size=9,
        color=(80, 80, 80)
    )
    return para


def add_english_line(cell, circle, eng_text):
    """Add English sentence with circle number."""
    para = cell.add_paragraph()
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.space_after = Pt(8)
    # Circle number
    add_run_with_style(para, circle + ' ', font_name='맑은 고딕', size=10, bold=True)
    # English text
    add_run_with_style(para, eng_text, font_name='Times New Roman', size=10)
    return para


def add_blank_line(cell, size=4):
    para = cell.add_paragraph()
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.space_after = Pt(0)
    run = para.add_run('')
    run.font.size = Pt(size)
    return para


def build_docx(pages, original_filename='output'):
    """Build a DOCX document from parsed pages."""
    doc = Document()

    # Page setup: A4
    section = doc.sections[0]
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.top_margin = Mm(15)
    section.bottom_margin = Mm(15)
    section.left_margin = Mm(16)
    section.right_margin = Mm(16)

    # Default paragraph style
    style = doc.styles['Normal']
    style.font.name = '맑은 고딕'
    style.font.size = Pt(10)

    for page_idx, page in enumerate(pages):
        if page_idx > 0:
            # Page break
            doc.add_page_break()

        # ── Header table ──────────────────────────────────
        hdr_tbl = doc.add_table(rows=1, cols=3)
        hdr_tbl.style = 'Table Grid'
        hdr_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

        # Col widths: logo | title | unit
        hdr_tbl.columns[0].width = Mm(35)
        hdr_tbl.columns[1].width = Mm(100)
        hdr_tbl.columns[2].width = Mm(55)

        c0, c1, c2 = hdr_tbl.rows[0].cells
        c0.width = Mm(35)
        c1.width = Mm(100)
        c2.width = Mm(55)

        p1 = c1.paragraphs[0]
        p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_run_with_style(p1, 'EBS 올림포스 – 영어독해 기본 01',
                           font_name='맑은 고딕', size=9, bold=True)

        p2 = c2.paragraphs[0]
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_run_with_style(p2, 'Unit 02 – 심경 (변화) · 분위기',
                           font_name='맑은 고딕', size=8)

        doc.add_paragraph('')  # spacing

        # ── Main content table ──────────────────────────────
        main_tbl = doc.add_table(rows=2, cols=1)
        set_table_border(main_tbl, border_type='single', color='000000', size=6)
        main_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

        # Row 0: Category header
        cat_cell = main_tbl.rows[0].cells[0]
        cat_p = cat_cell.paragraphs[0]
        cat_p.paragraph_format.space_before = Pt(2)
        cat_p.paragraph_format.space_after = Pt(2)
        cat_text = page.get('category', '▶ 함축 의미 추론')
        add_run_with_style(cat_p, cat_text, font_name='맑은 고딕', size=10, bold=True)

        # Row 1: Content
        content_cell = main_tbl.rows[1].cells[0]
        # Clear default empty paragraph
        content_cell._tc.clear_content()

        # Question paragraph
        if page.get('question_spans'):
            add_question_para(content_cell, page['question_spans'])
        elif page.get('question_raw'):
            q_p = content_cell.add_paragraph()
            add_run_with_style(q_p, page['question_raw'], size=10)

        add_blank_line(content_cell, 4)

        # Pairs: Korean above English
        if page.get('needs_transform', True):
            # Build pairs
            eng_dict = {circle: text for circle, text in page.get('english_pairs', [])}
            kor_dict = {circle: text for circle, text in page.get('korean_pairs', [])}
            circles = [c for c, _ in page.get('english_pairs', [])]

            for circle in circles:
                kor_text = kor_dict.get(circle, '')
                eng_text = eng_dict.get(circle, '')
                if kor_text:
                    add_korean_line(content_cell, kor_text)
                if eng_text:
                    add_english_line(content_cell, circle, eng_text)
                add_blank_line(content_cell, 3)
        else:
            # Already transformed: just use paired data
            for kor, circle, eng in page.get('paired', []):
                if kor:
                    add_korean_line(content_cell, kor)
                if eng:
                    add_english_line(content_cell, circle, eng)
                add_blank_line(content_cell, 3)

        # Table widths
        for row in main_tbl.rows:
            for cell in row.cells:
                cell.width = Mm(178)

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


# ─────────────────────────────────────────
# Routes
# ─────────────────────────────────────────

@app.route('/')
def index():
    return render_template('index.html')


@app.route('/transform', methods=['POST'])
def transform():
    if 'file' not in request.files:
        return jsonify({'error': '파일이 없습니다.'}), 400

    f = request.files['file']
    if not f.filename.lower().endswith('.hwp'):
        return jsonify({'error': 'HWP 파일만 지원합니다.'}), 400

    # Save uploaded file
    tmpdir = tempfile.mkdtemp()
    try:
        hwp_path = os.path.join(tmpdir, 'input.hwp')
        f.save(hwp_path)

        # Parse
        html = hwp_to_html(hwp_path)
        pages = parse_hwp_pages(html)

        if not pages:
            return jsonify({'error': '변환할 페이지를 찾을 수 없습니다. HWP 구조를 확인해주세요.'}), 400

        # Only transform pages that need it
        needs_transform = [p for p in pages if p.get('needs_transform')]
        stats = {
            'total_pages': len(pages),
            'transformed': len(needs_transform),
            'already_done': len(pages) - len(needs_transform),
        }

        # Build DOCX (all pages, but transformed ones converted)
        docx_buf = build_docx(pages)
        base_name = os.path.splitext(f.filename)[0]
        out_name = f'{base_name}_변환완료.docx'

        return send_file(
            docx_buf,
            as_attachment=True,
            download_name=out_name,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )

    except Exception as e:
        import traceback
        return jsonify({'error': str(e), 'trace': traceback.format_exc()}), 500
    finally:
        shutil.rmtree(tmpdir, ignore_errors=True)


@app.route('/preview', methods=['POST'])
def preview():
    """Return JSON preview of parsed pages for display."""
    if 'file' not in request.files:
        return jsonify({'error': '파일이 없습니다.'}), 400

    f = request.files['file']
    tmpdir = tempfile.mkdtemp()
    try:
        hwp_path = os.path.join(tmpdir, 'input.hwp')
        f.save(hwp_path)
        html = hwp_to_html(hwp_path)
        pages = parse_hwp_pages(html)

        # Summarize for preview
        preview_data = []
        for i, p in enumerate(pages):
            preview_data.append({
                'page': i + 1,
                'category': p.get('category', ''),
                'needs_transform': p.get('needs_transform', False),
                'sentence_count': len(p.get('english_pairs', p.get('paired', []))),
                'sample_english': (p.get('english_pairs') or p.get('paired') or [('', '')])[0],
                'sample_korean': (p.get('korean_pairs') or [('', '')])[0] if p.get('needs_transform') else None,
            })

        return jsonify({'pages': preview_data, 'total': len(pages)})
    except Exception as e:
        return jsonify({'error': str(e)}), 500
    finally:
        shutil.rmtree(tmpdir, ignore_errors=True)


if __name__ == '__main__':
    print("🚀 HWP 변환기 서버 시작: http://localhost:5000")
    app.run(debug=True, host='0.0.0.0', port=5000)
